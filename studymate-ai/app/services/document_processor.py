"""
services/document_processor.py
--------------------------------
Extracts text from uploaded documents.
Supports PDF, DOCX, TXT, CSV, and Images (via OCR/Vision).
"""

import io
from pathlib import Path
from typing import Optional


async def process_document(content: bytes, file_type: str, filename: str) -> Optional[str]:
    """
    Extract text from a document based on its type.
    Returns the extracted text string, or None if extraction fails.
    """
    try:
        if file_type == "pdf":
            return _extract_pdf(content)
        elif file_type == "docx":
            return _extract_docx(content)
        elif file_type == "txt":
            return _extract_txt(content)
        elif file_type == "csv":
            return _extract_csv(content)
        elif file_type == "image":
            return _extract_image_text(content)
        else:
            return None
    except Exception as e:
        return f"[Text extraction failed: {str(e)}]"


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages.append(f"--- Page {page_num + 1} ---\n{text}")
        doc.close()
        return "\n\n".join(pages) if pages else "[PDF has no readable text — may be scanned image]"
    except ImportError:
        return "[PyMuPDF not installed. Run: pip install PyMuPDF]"
    except Exception as e:
        return f"[PDF extraction error: {str(e)}]"


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs) if paragraphs else "[DOCX has no readable content]"
    except ImportError:
        return "[python-docx not installed. Run: pip install python-docx]"
    except Exception as e:
        return f"[DOCX extraction error: {str(e)}]"


def _extract_txt(content: bytes) -> str:
    """Extract text from plain text files."""
    for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_csv(content: bytes) -> str:
    """Extract text summary from CSV using pandas."""
    try:
        import pandas as pd
        df = pd.read_csv(io.BytesIO(content))

        lines = []
        lines.append(f"CSV File Summary:")
        lines.append(f"- Rows: {len(df)}")
        lines.append(f"- Columns: {len(df.columns)}")
        lines.append(f"- Column names: {', '.join(df.columns.tolist())}")
        lines.append("")
        lines.append("Data Types:")
        for col, dtype in df.dtypes.items():
            lines.append(f"  {col}: {dtype}")
        lines.append("")
        lines.append("Statistical Summary:")
        lines.append(df.describe(include="all").to_string())
        lines.append("")
        lines.append("First 10 rows:")
        lines.append(df.head(10).to_string())
        lines.append("")
        lines.append("Missing Values:")
        missing = df.isnull().sum()
        for col, count in missing[missing > 0].items():
            lines.append(f"  {col}: {count} missing")

        return "\n".join(lines)
    except ImportError:
        # Fallback: read as plain text
        return _extract_txt(content)
    except Exception as e:
        return f"[CSV extraction error: {str(e)}]\n\nRaw content:\n{_extract_txt(content)[:5000]}"


def _extract_image_text(content: bytes) -> str:
    """
    For images, return a description placeholder.
    Actual vision analysis is done via the AI provider's chat_with_image method.
    """
    try:
        # Try pytesseract OCR if available
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(img)
        if text.strip():
            return f"[OCR Extracted Text]\n\n{text}"
    except ImportError:
        pass
    except Exception:
        pass

    return "[Image uploaded — use the Image Analyzer to get AI analysis]"
